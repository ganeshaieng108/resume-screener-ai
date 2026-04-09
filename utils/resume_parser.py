"""
resume_parser.py
Extract clean text from PDF and DOCX resume files.
"""

import re
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_resume(file_path: str) -> str:
    """
    Dispatch to the correct parser based on file extension.
    Returns cleaned plain-text content of the resume.
    """
    path = Path(file_path)
    ext  = path.suffix.lower()

    if ext == ".pdf":
        text = _parse_pdf(file_path)
    elif ext == ".docx":
        text = _parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return _clean_text(text)


# ── PDF ───────────────────────────────────────────────────────────────────────

def _parse_pdf(file_path: str) -> str:
    """
    Try pdfplumber first (best for complex layouts), fall back to PyPDF2.
    """
    text = _parse_pdf_pdfplumber(file_path)
    if not text or len(text.split()) < 20:
        text = _parse_pdf_pypdf2(file_path)
    return text or ""


def _parse_pdf_pdfplumber(file_path: str) -> str:
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                raw = page.extract_text()
                if raw:
                    pages.append(raw)
        return "\n".join(pages)
    except Exception as e:
        logger.warning("pdfplumber failed: %s", e)
        return ""


def _parse_pdf_pypdf2(file_path: str) -> str:
    try:
        import PyPDF2
        pages = []
        with open(file_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                raw = page.extract_text()
                if raw:
                    pages.append(raw)
        return "\n".join(pages)
    except Exception as e:
        logger.warning("PyPDF2 failed: %s", e)
        return ""


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _parse_docx(file_path: str) -> str:
    """Extract text from all paragraphs and table cells in a .docx file."""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())

        return "\n".join(paragraphs)
    except Exception as e:
        logger.error("DOCX parsing failed: %s", e)
        return ""


# ── Text cleaning ─────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Normalise whitespace, remove non-printable characters, and strip
    common PDF artefacts while preserving meaningful punctuation.
    """
    if not text:
        return ""

    # Replace common ligatures / unicode quirks
    replacements = {
        "\u2019": "'", "\u2018": "'",
        "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "-",
        "\u00e2\u0080\u0099": "'",
        "\ufb01": "fi", "\ufb02": "fl",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)

    # Remove non-printable / control characters (keep newlines, tabs)
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E]", " ", text)

    # Collapse excessive whitespace but preserve paragraph breaks
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()
